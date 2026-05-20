"""Export legal analysis as DOCX. Optionally save to a case folder."""
import io
import re
import uuid
from typing import Optional, Literal

from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.auth import get_current_user_id
from app.database import get_supabase

router = APIRouter(prefix="/api/export", tags=["exports"])


class CaseMeta(BaseModel):
    title: Optional[str] = None
    case_number: Optional[str] = None
    case_type: Optional[str] = None
    court: Optional[str] = None
    plaintiff_name: Optional[str] = None
    defendant_name: Optional[str] = None
    claim_amount: Optional[float] = None
    perspective: Optional[str] = None


class SaveTo(BaseModel):
    case_id: str
    folder: Optional[str] = None


class AnalysisExportIn(BaseModel):
    analysis: str                                # markdown text
    case_meta: CaseMeta
    format: Literal["docx"] = "docx"
    filename: Optional[str] = None
    save_to: Optional[SaveTo] = None             # if set, upload to documents storage


# ── Markdown → DOCX renderer ──────────────────────────────────────────────────

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")
_INLINE_CODE_RE = re.compile(r"`([^`]+)`")


def _add_runs_with_inline_formatting(p, text: str):
    """Parse **bold**, *italic*, `code` and add styled runs to a paragraph."""
    # Tokenize by walking through the string, finding the earliest match each time
    pos = 0
    while pos < len(text):
        # Find earliest of bold / italic / code
        candidates = []
        m = _BOLD_RE.search(text, pos)
        if m:
            candidates.append(("bold", m))
        m = _ITALIC_RE.search(text, pos)
        if m:
            candidates.append(("italic", m))
        m = _INLINE_CODE_RE.search(text, pos)
        if m:
            candidates.append(("code", m))

        if not candidates:
            run = p.add_run(text[pos:])
            run.font.name = "Sarabun"
            break

        kind, m = min(candidates, key=lambda kv: kv[1].start())
        if m.start() > pos:
            run = p.add_run(text[pos:m.start()])
            run.font.name = "Sarabun"
        run = p.add_run(m.group(1))
        run.font.name = "Sarabun"
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        elif kind == "code":
            run.font.name = "Consolas"
        pos = m.end()


def _render_markdown_to_docx(doc: Document, md: str):
    """Render a (subset of) markdown into the given Document.

    Supports: # / ## / ### headings, bold **x**, italic *x*, code `x`,
    bullet lists (•, -, *), numbered lists (1. 2.), blank-line paragraphs,
    horizontal rule (---).
    """
    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$", stripped):
            p = doc.add_paragraph()
            p.add_run("―" * 30).font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
            i += 1
            continue

        # Headings
        if stripped.startswith("#### "):
            p = doc.add_heading(level=4)
            run = p.add_run(stripped[5:].strip())
            run.font.name = "Sarabun"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_heading(level=3)
            run = p.add_run(stripped[4:].strip())
            run.font.name = "Sarabun"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_heading(level=2)
            run = p.add_run(stripped[3:].strip())
            run.font.name = "Sarabun"
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_heading(level=1)
            run = p.add_run(stripped[2:].strip())
            run.font.name = "Sarabun"
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                txt = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Number")
                _add_runs_with_inline_formatting(p, txt)
                i += 1
            continue

        # Bullet list
        if re.match(r"^[•\-\*]\s+", stripped):
            while i < len(lines) and re.match(r"^[•\-\*]\s+", lines[i].strip()):
                txt = re.sub(r"^[•\-\*]\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Bullet")
                _add_runs_with_inline_formatting(p, txt)
                i += 1
            continue

        # Plain paragraph (consume until blank line)
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not _is_block_starter(lines[i].strip()):
            para_lines.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        _add_runs_with_inline_formatting(p, " ".join(para_lines))


def _is_block_starter(s: str) -> bool:
    if s.startswith("#"):
        return True
    if re.match(r"^\d+\.\s+", s):
        return True
    if re.match(r"^[•\-\*]\s+", s):
        return True
    if re.match(r"^-{3,}$", s):
        return True
    return False


_NAVY = RGBColor(0x1F, 0x2A, 0x44)


def _cover_para(doc: Document, text: str, *, size: int, bold: bool = False,
                space_before: int = 0, space_after: int = 0,
                color: RGBColor = _NAVY) -> None:
    """Add a centered cover paragraph with consistent Thai legal styling."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Sarabun"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def _build_docx(analysis_md: str, meta: CaseMeta) -> bytes:
    """Produce a polished .docx for a legal analysis report."""
    doc = Document()

    # Set default font for the whole doc
    style = doc.styles["Normal"]
    style.font.name = "Sarabun"
    style.font.size = Pt(11)

    # ── COVER PAGE (formal Thai legal style) ──────────────────────────────
    running = " — ".join(filter(None, [
        f"คดีหมายเลข {meta.case_number}" if meta.case_number else None,
        meta.court or None,
    ]))
    if running:
        _cover_para(doc, running, size=10, space_after=24)
    else:
        _cover_para(doc, "", size=10, space_after=24)

    # Main title
    _cover_para(doc, meta.title or "บทสรุปคดี",
                size=28, bold=True, space_before=24, space_after=12)
    _cover_para(doc, "บทสรุปคดีและแนวทางต่อสู้ขั้นสุด",
                size=15, bold=True, space_after=36)

    # Case number + court block
    if meta.case_number:
        _cover_para(doc, f"คดีหมายเลข {meta.case_number}",
                    size=13, bold=True, space_after=2)
    if meta.court:
        _cover_para(doc, meta.court, size=13, bold=True, space_after=24)

    # Charges / case type
    if meta.case_type:
        charges = f"ประเภทคดี: {meta.case_type}"
        if meta.claim_amount:
            charges += f"  ·  ทุนทรัพย์ ฿{meta.claim_amount:,.0f}"
        _cover_para(doc, charges, size=12, bold=True, space_after=20)

    # Parties block — "ระหว่าง" with role labels
    if meta.plaintiff_name or meta.defendant_name:
        rows = []
        if meta.plaintiff_name:
            rows.append(("ระหว่าง", meta.plaintiff_name, "โจทก์"))
        if meta.defendant_name:
            rows.append(("",        meta.defendant_name, "จำเลย"))

        table = doc.add_table(rows=len(rows), cols=3)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.autofit = True
        for r_idx, (label, name, role) in enumerate(rows):
            row = table.rows[r_idx]
            for cell, txt, bold in (
                (row.cells[0], label, True),
                (row.cells[1], name,  False),
                (row.cells[2], role,  True),
            ):
                cell_p = cell.paragraphs[0]
                cell_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = cell_p.add_run(txt)
                run.font.name = "Sarabun"
                run.font.size = Pt(12)
                run.bold = bold
                run.font.color.rgb = _NAVY

        _cover_para(doc, "", size=10, space_after=8)

    # Disclaimer at bottom
    _cover_para(doc, "เอกสารนี้เป็นแนวทางวิชาการประกอบการปรึกษาทนายความ",
                size=10, space_before=72)
    _cover_para(doc, "มิใช่คำปรึกษาทางกฎหมายอย่างเป็นทางการ", size=10)
    persp = meta.perspective or "ทั้งสองฝ่าย"
    from datetime import datetime
    today_th = datetime.now().strftime("%d/%m/%Y")
    _cover_para(doc, f"มุมมอง: {persp}  ·  จัดทำวันที่ {today_th}",
                size=10, space_after=24)
    _cover_para(doc, "— Thai.Law · Legal Memo —",
                size=9, bold=True, space_after=0)

    doc.add_page_break()

    # Body content
    _render_markdown_to_docx(doc, analysis_md)

    # Serialize
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _safe_filename(name: str) -> str:
    name = (name or "analysis").strip()
    # Strip anything that would break a filename
    name = re.sub(r"[\\/:*?\"<>|\r\n\t]+", " ", name)
    name = name[:80].strip() or "analysis"
    return name


# ── Endpoint ──────────────────────────────────────────────────────────────────

@router.post("/analysis")
async def export_analysis(
    payload: AnalysisExportIn,
    user_id: str = Depends(get_current_user_id),
):
    if not payload.analysis or not payload.analysis.strip():
        raise HTTPException(status_code=400, detail="analysis is empty")

    blob = _build_docx(payload.analysis, payload.case_meta)
    base_name = _safe_filename(payload.filename or payload.case_meta.title or "analysis")
    file_name = f"{base_name}.docx"

    # If save_to is provided → upload to Supabase Storage and insert doc row
    if payload.save_to:
        db = get_supabase()
        storage_path = f"{user_id}/{uuid.uuid4()}.docx"
        try:
            db.storage.from_("documents").upload(
                path=storage_path,
                file=blob,
                file_options={
                    "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "upsert": "false",
                },
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Storage upload failed: {str(e)}")

        row = {
            "user_id":       user_id,
            "case_id":       payload.save_to.case_id,
            "title":         file_name,
            "doc_label":     base_name,
            "original_name": file_name,
            "doc_category":  "บทวิเคราะห์",
            "file_type":     "docx",
            "file_size":     len(blob),
            "storage_path":  storage_path,
            "is_processed":  True,
            "ai_summary":    "บทสรุปคดีและแนวทางต่อสู้ขั้นสุด (Thai.Law Legal Memo)",
            "folder":        (payload.save_to.folder or None),
        }
        try:
            response = db.table("documents").insert(row).execute()
            if not response.data:
                raise HTTPException(status_code=500, detail="DB insert returned empty")
            return {"saved": True, "document": response.data[0]}
        except HTTPException:
            raise
        except Exception as e:
            try: db.storage.from_("documents").remove([storage_path])
            except Exception: pass
            raise HTTPException(status_code=500, detail=f"DB insert failed: {str(e)}")

    # Otherwise stream as download
    return StreamingResponse(
        io.BytesIO(blob),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{file_name}"',
            "Content-Length": str(len(blob)),
        },
    )
