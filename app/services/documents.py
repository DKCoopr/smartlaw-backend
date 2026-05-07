"""Document AI: extract summary from PDF/image/DOCX/DOC using Gemini."""
import io
import shutil
import subprocess
import tempfile
import google.generativeai as genai
from app.config import get_settings

settings = get_settings()
genai.configure(api_key=settings.google_api_key)

# Pro model handles Thai legal documents more accurately than Flash.
_model = genai.GenerativeModel("gemini-1.5-pro-latest")

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_MIME  = "application/msword"


def _extract_doc_text(file_bytes: bytes) -> str:
    """Pull text out of a legacy .doc (Word binary) file. Tries antiword
    first, then catdoc as fallback. Both are installed via Dockerfile.
    Returns "" on any failure (best-effort, never raises)."""
    if not file_bytes:
        return ""
    # Write bytes to a temp file so the CLI tools can read by path
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
    except Exception:
        return ""

    try:
        for tool in ("antiword", "catdoc"):
            if not shutil.which(tool):
                continue
            try:
                # Both tools print plain text to stdout when given a path
                result = subprocess.run(
                    [tool, tmp_path],
                    capture_output=True,
                    timeout=30,
                )
                if result.returncode == 0:
                    text = (result.stdout or b"").decode("utf-8", errors="replace").strip()
                    if text:
                        return text
            except Exception:
                continue
        return ""
    finally:
        try:
            import os
            os.unlink(tmp_path)
        except Exception:
            pass


def _extract_docx_text(file_bytes: bytes) -> str:
    """Pull plain text out of a .docx so Gemini (which can't ingest DOCX
    inline) can summarize it as text/plain. Best-effort: returns "" on any
    parse failure."""
    try:
        from docx import Document   # python-docx — already in requirements.txt
    except Exception:
        return ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        # Tables — flatten cell-by-cell so structured docs (contracts/letters)
        # don't lose all numeric/tabular content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join((c.text or "").strip() for c in row.cells if (c.text or "").strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts).strip()
    except Exception:
        return ""

SUMMARY_PROMPT = """คุณเป็นทนายความไทยผู้เชี่ยวชาญ อ่านเอกสารนี้ (เป็นภาษาไทย) แล้วเขียนสรุปเชิงกฎหมายอย่างละเอียดในภาษาไทยทางการ

โครงสร้างคำตอบ:

**ประเภทเอกสาร:** (เช่น คำฟ้อง / คำให้การ / สัญญา / หนังสือบอกกล่าว / พินัยกรรม / คำสั่งศาล / รายงานแพทย์ / ใบรับรองแพทย์ / ใบเสร็จ ฯลฯ)

**คู่ความ/ผู้เกี่ยวข้อง:**
- ชื่อ-นามสกุลของบุคคลและนิติบุคคลทุกฝ่ายที่ปรากฏในเอกสาร (เขียนชื่อจริง)
- บทบาทของแต่ละฝ่าย (โจทก์/จำเลย/ผู้ให้/ผู้รับ ฯลฯ)

**สาระสำคัญ:**
- ใจความหลักของเอกสาร 3-5 บรรทัด
- วันที่/เลขเอกสาร/เลขคดีที่ปรากฏ (ถ้ามี)
- จำนวนเงิน/ทรัพย์สินที่ระบุ (ถ้ามี)

**ข้อกฎหมายที่อ้าง:**
- มาตรา / พระราชบัญญัติ / คำพิพากษาฎีกา ที่เอกสารกล่าวถึง

**จุดสำคัญ/ความเสี่ยงทางกฎหมาย:**
- ประเด็นที่ทนายต้องสนใจเป็นพิเศษ (เงื่อนไขเสี่ยง / ข้อสัญญาผิดปกติ / ข้อมูลที่ขาดหาย / ความขัดแย้ง)

ตอบเฉพาะเนื้อหาที่ปรากฏจริงในเอกสาร ห้ามแต่งเติม ถ้าข้อมูลส่วนใดไม่มีให้ใส่ "—" """

# Mime types Gemini can read inline
SUPPORTED_INLINE_MIMES = {
    "application/pdf",
    "image/png",
    "image/jpeg",
    "image/webp",
    "image/heic",
    "image/heif",
    "text/plain",
}

MAX_INLINE_BYTES = 18 * 1024 * 1024   # 18MB safety margin under Gemini 20MB limit


async def summarize_document(file_bytes: bytes, mime_type: str) -> str:
    """
    Returns a Thai-language summary of the document, or "" if unsupported/too large.
    Errors are swallowed and return "" — summary is best-effort.

    DOCX is handled separately: python-docx extracts the text first, then we
    feed it to Gemini as text/plain (Gemini cannot ingest DOCX inline).
    """
    if len(file_bytes) > MAX_INLINE_BYTES:
        return ""

    # DOCX / DOC path: extract text → call Gemini as text/plain
    # (Gemini cannot ingest these formats inline)
    if mime_type in (DOCX_MIME, DOC_MIME):
        text = _extract_docx_text(file_bytes) if mime_type == DOCX_MIME else _extract_doc_text(file_bytes)
        if not text:
            return ""
        try:
            response = await _model.generate_content_async([
                {"mime_type": "text/plain", "data": text.encode("utf-8")},
                SUMMARY_PROMPT,
            ])
            return (response.text or "").strip()
        except Exception:
            return ""

    if mime_type not in SUPPORTED_INLINE_MIMES:
        return ""
    try:
        response = await _model.generate_content_async([
            {"mime_type": mime_type, "data": file_bytes},
            SUMMARY_PROMPT,
        ])
        return (response.text or "").strip()
    except Exception as e:
        return f""


# ── OCR-only mode: extract verbatim text without analyzing ───────────────────

OCR_PROMPT = """ทำหน้าที่เป็น OCR เท่านั้น — อ่านข้อความทั้งหมดในภาพให้ครบ ห้ามวิเคราะห์ ห้ามสรุป ห้ามตัดทอน

กฎ:
1. ดึงข้อความทุกตัวอักษรในภาพ — ทั้งพิมพ์และเขียนด้วยลายมือ
2. คงรูปแบบเดิม — ย่อหน้า เลขข้อ ตัวเอน ตัวหนา ตำแหน่ง บรรทัด
3. ถ้ามีตาราง — เขียนเป็น markdown table
4. ถ้าอ่านไม่ออกบางส่วน — ใส่ [อ่านไม่ออก] ตรงนั้น แทนการเดา
5. ถ้ามีตัวเลข วันที่ ลายเซ็น ตราประทับ — บันทึกตามที่เห็นทุกตัว
6. ถ้าหลายภาพต่อเนื่องกัน — แยกแต่ละหน้าด้วย "--- หน้าถัดไป ---"

ตอบเฉพาะข้อความที่อ่านได้จากภาพ ห้ามมีคำอธิบายอื่น ห้ามมีคำเริ่ม "นี่คือ..." หรือ "ข้อความในภาพคือ..."""


async def ocr_image(file_bytes: bytes, mime_type: str) -> str:
    """
    Extract verbatim text from image (handwriting + printed). No analysis.
    Used to feed Claude with full text content from scanned/photographed documents.
    """
    if not mime_type.startswith("image/"):
        return ""
    if len(file_bytes) > MAX_INLINE_BYTES:
        return ""
    try:
        response = await _model.generate_content_async([
            {"mime_type": mime_type, "data": file_bytes},
            OCR_PROMPT,
        ])
        return (response.text or "").strip()
    except Exception:
        return ""
